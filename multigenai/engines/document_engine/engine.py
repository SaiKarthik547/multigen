"""
DocumentEngine — Word document generation using python-docx + Wikipedia content.

Phase 1: Migrated from MultiGenAi.py with ExecutionContext injection and
         typed request/response.
Phase 7: Will add LLM-driven content planning, style-based theming,
         automatic image insertion, and layout consistency.
"""

from __future__ import annotations

import pathlib
import re
from dataclasses import dataclass
from typing import TYPE_CHECKING, Optional

from multigenai.core.logging.logger import get_logger

if TYPE_CHECKING:
    from multigenai.core.execution_context import ExecutionContext
    from multigenai.llm.schema_validator import DocumentGenerationRequest

LOG = get_logger(__name__)


@dataclass
class DocumentResult:
    """Output from the DocumentEngine."""
    path: str
    title: str
    page_count_estimate: int
    success: bool = True
    error: Optional[str] = None


class DocumentEngine:
    """
    Generates Word (.docx) documents from a natural language prompt.

    Usage:
        engine = DocumentEngine(ctx)
        result = engine.run(DocumentGenerationRequest(prompt="quantum computing", output_format="docx"))
    """

    def __init__(self, ctx: "ExecutionContext") -> None:
        self._ctx = ctx
        self._out_dir = pathlib.Path(ctx.settings.output_dir)
        self._out_dir.mkdir(parents=True, exist_ok=True)

    def run(self, request: "DocumentGenerationRequest") -> DocumentResult:
        """Generate a document from the validated request."""
        import hashlib
        slug = re.sub(r"[^A-Za-z0-9\-_]+", "_", request.prompt)[:40]
        slug += f"_{hashlib.sha1(request.prompt.encode()).hexdigest()[:8]}"
        out_path = self._out_dir / f"{slug}.{request.output_format}"

        title, content = self._fetch_content(request.prompt)

        try:
            import docx as python_docx
        except ImportError:
            msg = "python-docx not installed — run: pip install python-docx"
            LOG.error(msg)
            return DocumentResult(path="", title=title, page_count_estimate=0, success=False, error=msg)

        document = python_docx.Document()
        document.add_heading(title, level=0)
        document.add_paragraph(f"Generated document about: {request.prompt}")
        if content:
            for para in content.split("\n"):
                if para.strip():
                    document.add_paragraph(para)
        document.save(str(out_path))
        LOG.info(f"Document saved: {out_path}")
        paragraphs = len([p for p in content.split("\n") if p.strip()])
        return DocumentResult(path=str(out_path), title=title, page_count_estimate=max(1, paragraphs // 4))

    # ------------------------------------------------------------------
    # Content sourcing
    # ------------------------------------------------------------------

    def _fetch_content(self, query: str) -> tuple[str, str]:
        """Fetch content from Wikipedia, graceful fallback if unavailable."""
        try:
            import wikipediaapi
            wiki = wikipediaapi.Wikipedia("MultiGenAI/1.0", "en")
            page = wiki.page(query)
            if page.exists():
                summary = " ".join(page.summary.split()[:400]) + "…"
                return page.title, summary
        except ImportError:
            LOG.warning("wikipedia-api not installed — using placeholder content.")
        except Exception as exc:
            LOG.warning(f"Wikipedia fetch failed: {exc}")
        return f"Document: {query}", f"This document covers the topic: {query}."

    # Phase 7 hook
    def run_with_llm_planning(self, request: "DocumentGenerationRequest") -> DocumentResult:
        """[Phase 7] LLM-driven structured content planning with style theming."""
        raise NotImplementedError("LLM document planning activates in Phase 7.")
